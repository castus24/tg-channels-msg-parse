import datetime
import telethon
import asyncio
from docx import Document


async def main():
    # Telegram API keys
    api_id = #API_ID
    api_hash = #'API_HASH'

    # Channel name or username
    channel_name = #'channel_name'

    # Current date
    today = datetime.datetime.now().date()

    # Create a new Telegram client
    client = telethon.TelegramClient('session_name', api_id, api_hash)

    # Connect to Telegram
    await client.start()

    # Ensure the connection is still alive
    if not client.is_connected():
        await client.connect()

    # Get the channel entity
    channel = await client.get_entity(channel_name)

    # Create a new word document
    document = Document()

    # Get all the messages from the channel
    messages = [msg async for msg in client.iter_messages(channel)]

    # Filter the messages by date
    today_messages = [msg for msg in messages if msg.date.date() == today]

    # Iterate through the messages and add them to the word document without images
    # for message in today_messages:
    #     document.add_paragraph(message.message)

    # Iterate through the messages and add them to the word document with images
    for message in today_messages:
        if message.media:
            if message.photo:
                # Download the image
                file_name = f'photo_{message.photo.id}.jpg'
                await message.download_media(file=file_name)
                print(f'Image saved: {file_name}')
                # Add the image to the word document
                document.add_picture(file_name)
                document.add_paragraph(message.message)
            elif message.document and message.document.mime_type.startswith('image/'):
                file_name = f'document_{message.document.id}.{message.document.mime_type.split("/")[-1]}'
                await message.download_media(file=file_name)
                print(f'Image saved: {file_name}')
                # Add the image to the word document
                document.add_picture(file_name)
                document.add_paragraph(message.message)
        else:
            document.add_paragraph(message.message)

    # Save the word document
    document.save('messages.docx')
    # Log out
    await client.log_out()


if __name__ == '__main__':
    asyncio.run(main())